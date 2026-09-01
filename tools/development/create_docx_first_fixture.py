from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


BLUE = "2F5597"
DARK_BLUE = "17365D"
LIGHT_BLUE = "EAF0F8"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D9E2F3"
WHITE = "FFFFFF"
TEXT = "202124"
MUTED = "666666"
FONT = "Malgun Gothic"
CONTENT_DXA = 9865  # A4 width 210 mm - 18 mm margins on both sides.


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    if sum(widths) != CONTENT_DXA:
        raise ValueError(f"table widths must total {CONTENT_DXA}: {widths}")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(index, len(widths) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_run(run, size=10, bold=False, color=TEXT, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph(paragraph, before=0, after=0, line=1.0, align=None):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def write_cell(cell, text, *, bold=False, size=9.5, color=TEXT, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_paragraph(paragraph, line=1.0, align=align)
    set_run(paragraph.add_run(text), size=size, bold=bold, color=color)


def style_label(cell, text):
    set_cell_shading(cell, LIGHT_BLUE)
    write_cell(cell, text, bold=True, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)


def add_section_label(doc, text):
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, before=4, after=3, line=1.0)
    set_run(paragraph.add_run(text), size=11, bold=True, color=DARK_BLUE)


def add_header_footer(section):
    header = section.header
    hp = header.paragraphs[0]
    set_paragraph(hp, after=0, line=1.0)
    set_run(hp.add_run("DOCX → HWPX 호환성 시험"), size=8.5, bold=True, color=MUTED)
    hp.add_run("\t")
    set_run(hp.add_run("DocBridge"), size=8.5, color=MUTED)
    hp.paragraph_format.tab_stops.add_tab_stop(Mm(174))

    footer = section.footer
    fp = footer.paragraphs[0]
    set_paragraph(fp, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run(fp.add_run("검증용 문서  |  A4 1쪽  |  1 / 1"), size=8, color=MUTED)


def create_document(output_path):
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(14)
    section.bottom_margin = Mm(14)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)
    section.header_distance = Mm(7)
    section.footer_distance = Mm(7)
    add_header_footer(section)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.10

    title = doc.add_paragraph()
    set_paragraph(title, before=3, after=2, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run(title.add_run("현장기술인 변경계"), size=20, bold=True, color=DARK_BLUE)
    subtitle = doc.add_paragraph()
    set_paragraph(subtitle, after=7, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run(subtitle.add_run("DOCX-first 한글 변환 호환성·속도 시험 문서"), size=9, color=MUTED)

    info = doc.add_table(rows=3, cols=4)
    info.style = "Table Grid"
    set_table_geometry(info, [1350, 3582, 1350, 3583])
    info_values = [
        ("공 사 명", "성동구 하수관로 정비공사", "문서번호", "DB-HWP-2026-0813"),
        ("현 장 명", "성수처리분구", "작성일", "2026-08-13"),
        ("발 주 처", "성동구청", "시공사", "청도건설(주)"),
    ]
    for row, values in zip(info.rows, info_values):
        prevent_row_split(row)
        for index, value in enumerate(values):
            if index % 2 == 0:
                style_label(row.cells[index], value)
            else:
                write_cell(row.cells[index], value)

    add_section_label(doc, "1. 변경 대상")
    people = doc.add_table(rows=3, cols=5)
    people.style = "Table Grid"
    set_table_geometry(people, [1180, 2050, 1500, 2785, 2350])
    headers = ["구분", "성명", "직급", "자격·등급", "배치일"]
    for index, text in enumerate(headers):
        set_cell_shading(people.rows[0].cells[index], BLUE)
        write_cell(people.rows[0].cells[index], text, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_table_header(people.rows[0])
    person_rows = [
        ["변경 전", "김현장", "차장", "토목기사 / 중급", "2025-03-01"],
        ["변경 후", "홍길동", "차장", "토목기사 / 고급", "2026-08-17"],
    ]
    for row, values in zip(people.rows[1:], person_rows):
        prevent_row_split(row)
        for index, value in enumerate(values):
            write_cell(row.cells[index], value, align=WD_ALIGN_PARAGRAPH.CENTER)
        if values[0] == "변경 후":
            set_cell_shading(row.cells[0], MID_GRAY)

    add_section_label(doc, "2. 변경 사유 및 업무 인계")
    reason = doc.add_table(rows=2, cols=2)
    reason.style = "Table Grid"
    set_table_geometry(reason, [1650, 8215])
    reason_values = [
        ("변경 사유", "기존 현장기술인의 본사 전보에 따라 공정 연속성과 품질·안전관리 강화를 위해 교체 배치함."),
        ("인계 사항", "설계도서, 공정표, 품질시험 성과, 안전점검 기록 및 관계기관 협의 현황을 2026-08-16까지 서면 인계함."),
    ]
    for row, values in zip(reason.rows, reason_values):
        prevent_row_split(row)
        style_label(row.cells[0], values[0])
        write_cell(row.cells[1], values[1], size=9.2)

    add_section_label(doc, "3. 첨부 및 확인")
    check = doc.add_table(rows=3, cols=4)
    check.style = "Table Grid"
    set_table_geometry(check, [2450, 2482, 2450, 2483])
    check_values = [
        ("■ 재직증명서", "1부", "■ 기술자 경력증명서", "1부"),
        ("■ 자격증 사본", "1부", "■ 배치계획서", "1부"),
        ("■ 4대보험 가입확인", "1부", "□ 기타", "해당 없음"),
    ]
    for row, values in zip(check.rows, check_values):
        prevent_row_split(row)
        for index, value in enumerate(values):
            if index % 2 == 0:
                set_cell_shading(row.cells[index], LIGHT_GRAY)
                write_cell(row.cells[index], value, bold=True, color=DARK_BLUE)
            else:
                write_cell(row.cells[index], value, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_section_label(doc, "4. 변경 후 관리계획")
    management = doc.add_table(rows=4, cols=2)
    management.style = "Table Grid"
    set_table_geometry(management, [1650, 8215])
    management_values = [
        ("업무 개시", "발주처 승인 즉시 현장기술인 배치계와 관계기관 신고사항을 갱신하고 공정별 책임 범위를 확인함."),
        ("안전관리", "매일 작업 전 TBM을 주관하고 위험성평가 조치사항 및 건설기계 작업계획의 현장 이행 여부를 점검함."),
        ("품질관리", "자재 승인, 품질시험, 시공검측 기록을 주간 단위로 대조하고 미결사항을 다음 공정 전에 종결함."),
        ("보고체계", "중대 위험·품질 부적합은 즉시 현장소장에게 보고하고 조치 결과를 발주처 정기회의 자료에 반영함."),
    ]
    for row, values in zip(management.rows, management_values):
        prevent_row_split(row)
        style_label(row.cells[0], values[0])
        write_cell(row.cells[1], values[1], size=8.8)

    certification = doc.add_paragraph()
    set_paragraph(certification, before=7, after=7, line=1.1, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run(certification.add_run("위와 같이 현장기술인 변경을 제출하오니 검토하여 주시기 바랍니다."), size=10.5, bold=True)

    approval = doc.add_table(rows=2, cols=4)
    approval.style = "Table Grid"
    set_table_geometry(approval, [1650, 3282, 1650, 3283])
    approval_values = [
        ("작성", "현장대리인  홍길동  (서명)", "검토", "품질관리자  이수원  (서명)"),
        ("승인", "현장소장  박성호  (서명)", "접수", "발주처 담당자              (서명)"),
    ]
    for row, values in zip(approval.rows, approval_values):
        prevent_row_split(row)
        for index, value in enumerate(values):
            if index % 2 == 0:
                set_cell_shading(row.cells[index], BLUE)
                write_cell(row.cells[index], value, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                write_cell(row.cells[index], value, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Leave the final table as the last body element. An explicit empty
    # paragraph after the table is harmless in Word, but Hancom's OOXML
    # importer can count it as a blank second page.

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


if __name__ == "__main__":
    if len(sys.argv) != 2:
        raise SystemExit("usage: create_docx_first_fixture.py OUTPUT.docx")
    create_document(Path(sys.argv[1]).resolve())
